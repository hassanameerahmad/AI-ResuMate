import io
import json
import re
from typing import Any

import streamlit as st
from pypdf import PdfReader
from docx import Document
from google import genai
from google.genai import types

st.set_page_config(page_title="Resume ATS Analyzer", page_icon="📄", layout="wide")

MODEL = "gemini-3.5-flash"
MAX_TEXT_CHARS = 30000


def extract_text(uploaded_file) -> str:
    """Extract readable text from PDF, DOCX, or TXT uploads."""
    file_type = uploaded_file.name.lower().rsplit(".", 1)[-1]

    if file_type == "pdf":
        reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif file_type == "docx":
        document = Document(io.BytesIO(uploaded_file.getvalue()))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(parts)
    elif file_type == "txt":
        text = uploaded_file.getvalue().decode("utf-8", errors="replace")
    else:
        raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")

    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        raise ValueError(
            "No selectable text was found. If this is a scanned/image-only PDF, "
            "please use a text-based PDF or DOCX."
        )
    return text[:MAX_TEXT_CHARS]


def heuristic_ats_checks(resume_text: str) -> dict[str, Any]:
    """Basic deterministic checks to complement the AI assessment."""
    text = resume_text.lower()
    checks = {
        "contact_info": bool(re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", resume_text))
        and bool(re.search(r"(?:\+?\d[\d\s().-]{7,}\d)", resume_text)),
        "common_sections": sum(
            section in text
            for section in ["experience", "education", "skills", "projects"]
        )
        >= 3,
        "action_verbs": len(
            re.findall(
                r"\b(achieved|built|created|developed|designed|improved|implemented|led|managed|optimized|reduced|increased|automated)\b",
                text,
            )
        )
        >= 3,
        "quantified_results": bool(
            re.search(r"\b\d+(?:\.\d+)?\s*(?:%|percent|x|k|m|million|thousand)\b", text)
        ),
        "reasonable_length": 250 <= len(resume_text.split()) <= 1500,
        "keyword_density": len(re.findall(r"\b[a-zA-Z][a-zA-Z+#.-]{2,}\b", resume_text)) >= 120,
    }
    score = round(sum(checks.values()) / len(checks) * 100)
    return {"score": score, "checks": checks}


def analyze_with_gemini(resume_text: str, job_description: str) -> dict[str, Any]:
    api_key = st.secrets.get("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError(
            "GEMINI_API_KEY is missing. Add it to Streamlit Secrets or set it as an environment variable."
        )

    client = genai.Client(api_key=api_key)

    schema = {
        "type": "object",
        "properties": {
            "ats_score": {"type": "integer", "description": "Overall ATS compatibility score from 0 to 100."},
            "summary": {"type": "string"},
            "strengths": {"type": "array", "items": {"type": "string"}},
            "critical_issues": {"type": "array", "items": {"type": "string"}},
            "improvements": {"type": "array", "items": {"type": "string"}},
            "missing_keywords": {"type": "array", "items": {"type": "string"}},
            "formatting_advice": {"type": "array", "items": {"type": "string"}},
            "rewritten_bullets": {"type": "array", "items": {"type": "string"}},
        },
        "required": [
            "ats_score",
            "summary",
            "strengths",
            "critical_issues",
            "improvements",
            "missing_keywords",
            "formatting_advice",
            "rewritten_bullets",
        ],
    }

    job_context = job_description.strip() if job_description.strip() else "No job description supplied. Evaluate general ATS readiness and do not invent job-specific keywords."

    prompt = f"""
You are an expert ATS resume reviewer and technical recruiter.
Analyze the resume below. Give a practical ATS compatibility score from 0-100.
This score is an estimate, not a score from a specific ATS vendor.

Scoring guidance:
- 25 points: clear standard sections and ATS-readable structure
- 25 points: relevant keywords and skills, especially against the job description
- 20 points: strong experience/project bullets with action verbs and measurable outcomes
- 15 points: completeness (contact info, education, skills, experience/projects)
- 15 points: concise, professional, ATS-friendly formatting

Rules:
1. Never invent experience, education, certifications, employers, or skills.
2. Identify missing keywords only when they are clearly supported by the job description or are standard terms relevant to it.
3. Prefer plain headings, standard terminology, measurable achievements, and simple formatting.
4. Keep recommendations specific and actionable.
5. Rewritten bullets must be based only on information actually present in the resume; if a metric is missing, use a placeholder such as [X%] rather than inventing one.

JOB DESCRIPTION:
{job_context}

RESUME:
{resume_text}
"""

    response = client.models.generate_content(
        model=MODEL,
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=schema,
            temperature=0.2,
        ),
    )

    if not response.text:
        raise RuntimeError("Gemini returned an empty response.")
    result = json.loads(response.text)
    result["ats_score"] = max(0, min(100, int(result["ats_score"])))
    return result


def display_list(items: list[str], empty_message: str = "None identified.") -> None:
    if not items:
        st.info(empty_message)
        return
    for item in items:
        st.markdown(f"- {item}")


st.title("📄 Resume ATS Analyzer")
st.caption("Upload your resume → get an estimated ATS score → see exactly what to improve.")

with st.sidebar:
    st.header("Settings")
    st.markdown("**AI model:** Gemini 2.5 Flash")
    st.markdown("**Supported:** PDF, DOCX, TXT")
    st.info("Tip: Add the target job description for a more useful keyword match.")

uploaded_file = st.file_uploader(
    "Upload your resume",
    type=["pdf", "docx", "txt"],
    help="Use a text-based PDF/DOCX. Scanned image-only PDFs are not OCR'd by this version.",
)

job_description = st.text_area(
    "Target job description (optional, but recommended)",
    height=220,
    placeholder="Paste the job description here to get job-specific keyword matching and recommendations...",
)

analyze = st.button("🔍 Analyze Resume", type="primary", use_container_width=True)

if analyze:
    if uploaded_file is None:
        st.warning("Please upload a resume first.")
        st.stop()

    try:
        resume_text = extract_text(uploaded_file)
    except Exception as exc:
        st.error(f"Could not read the resume: {exc}")
        st.stop()

    with st.spinner("Analyzing your resume with Gemini Flash..."):
        try:
            ai_result = analyze_with_gemini(resume_text, job_description)
        except Exception as exc:
            st.error(f"AI analysis failed: {exc}")
            st.stop()

    heuristic = heuristic_ats_checks(resume_text)
    ai_score = ai_result["ats_score"]

    st.session_state["analysis"] = ai_result
    st.session_state["heuristic"] = heuristic
    st.session_state["resume_name"] = uploaded_file.name

if "analysis" in st.session_state:
    result = st.session_state["analysis"]
    heuristic = st.session_state["heuristic"]

    st.divider()
    st.subheader(f"Results — {st.session_state.get('resume_name', 'Resume')}")

    col1, col2 = st.columns(2)
    with col1:
        st.metric("Estimated ATS Score", f"{result['ats_score']}/100")
        st.progress(result["ats_score"] / 100)
    with col2:
        st.metric("Basic ATS Readability Checks", f"{heuristic['score']}/100")
        st.progress(heuristic["score"] / 100)

    st.caption("The ATS score is an AI-based estimate. Different ATS products and employers use different parsing and ranking rules.")

    st.subheader("📝 Overall assessment")
    st.write(result["summary"])

    tab1, tab2, tab3, tab4 = st.tabs(["🚨 Critical Issues", "💪 Strengths", "🎯 Keywords", "✍️ Improvements"])

    with tab1:
        display_list(result["critical_issues"])

    with tab2:
        display_list(result["strengths"])

    with tab3:
        display_list(result["missing_keywords"], "No major missing keywords were identified from the supplied job description.")

    with tab4:
        st.markdown("**Priority improvements**")
        display_list(result["improvements"])
        st.markdown("**Formatting / ATS readability**")
        display_list(result["formatting_advice"])
        st.markdown("**Example rewritten bullets**")
        display_list(result["rewritten_bullets"])

    with st.expander("Basic deterministic checks"):
        labels = {
            "contact_info": "Email + phone detected",
            "common_sections": "Common resume sections detected",
            "action_verbs": "Action verbs used",
            "quantified_results": "At least one quantified result detected",
            "reasonable_length": "Resume length is within a broad ATS-friendly range",
            "keyword_density": "Enough text was extracted for keyword analysis",
        }
        for key, passed in heuristic["checks"].items():
            st.write(("✅" if passed else "⚠️") + " " + labels[key])

    st.download_button(
        "⬇️ Download analysis as JSON",
        data=json.dumps(result, indent=2, ensure_ascii=False),
        file_name="resume_ats_analysis.json",
        mime="application/json",
    )
